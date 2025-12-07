import os
import io
import fitz
from docx import Document
from PIL import Image
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# Tentar importar pytesseract (opcional)
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    print("⚠️ pytesseract não instalado. Análise de texto em imagens será limitada.")
def extract_text_from_pdf(file_content):
    """
    Extrai texto de PDF com estrutura preservada
    """
    try:
        pdf_doc = fitz.open(stream=file_content, filetype="pdf")
        extracted_data = {
            "text_content": "",
            "structure": {
                "pages": [],
                "headings": [],
                "paragraphs": [],
                "tables": []
            },
            "metadata": {
                "total_pages": len(pdf_doc),
                "title": pdf_doc.metadata.get("title", ""),
                "author": pdf_doc.metadata.get("author", "")
            }
        }
        
        full_text = ""
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            
            blocks = page.get_text("dict")
            page_text = ""
            page_structure = {
                "page_number": page_num + 1,
                "headings": [],
                "paragraphs": [],
                "images": []
            }
            
            for block in blocks["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_text = ""
                        font_size = 0
                        
                        for span in line["spans"]:
                            line_text += span["text"]
                            font_size = max(font_size, span["size"])
                        
                        if line_text.strip():
                            if font_size > 14:
                                page_structure["headings"].append({
                                    "text": line_text.strip(),
                                    "level": 1 if font_size > 18 else 2,
                                    "page": page_num + 1
                                })
                                extracted_data["structure"]["headings"].append({
                                    "text": line_text.strip(),
                                    "level": 1 if font_size > 18 else 2,
                                    "page": page_num + 1
                                })
                            else:
                                page_structure["paragraphs"].append(line_text.strip())
                                extracted_data["structure"]["paragraphs"].append({
                                    "text": line_text.strip(),
                                    "page": page_num + 1
                                })
                            
                            page_text += line_text + "\n"
            
            extracted_data["structure"]["pages"].append(page_structure)
            full_text += f"\n--- Página {page_num + 1} ---\n" + page_text
        
        pdf_doc.close()
        extracted_data["text_content"] = full_text
        return extracted_data
        
    except Exception as e:
        return {"erro": f"Erro ao processar PDF: {str(e)}"}

def extract_text_from_docx(file_content):
    """
    Extrai texto de DOCX preservando estrutura
    """
    try:
        doc = Document(io.BytesIO(file_content))
        
        extracted_data = {
            "text_content": "",
            "structure": {
                "headings": [],
                "paragraphs": [],
                "tables": []
            }
        }
        
        full_text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                    extracted_data["structure"]["headings"].append({
                        "text": para.text.strip(),
                        "level": level
                    })
                else:
                    extracted_data["structure"]["paragraphs"].append({
                        "text": para.text.strip()
                    })
                
                full_text += para.text + "\n"
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            extracted_data["structure"]["tables"].append({
                "data": table_data,
                "rows": len(table_data),
                "columns": len(table_data[0]) if table_data else 0
            })
        
        extracted_data["text_content"] = full_text
        return extracted_data
        
    except Exception as e:
        return {"erro": f"Erro ao processar DOCX: {str(e)}"}



def generate_document_summary(text_content):
    """
    Gera resumo automático do documento usando Gemini
    """
    try:
        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        model = genai.GenerativeModel(model_name="gemini-2.5-flash")
        
        prompt = f"""
        Analise o seguinte texto e gere um resumo estruturado e acessível:
        
        1. Resumo principal (2-3 frases)
        2. Pontos-chave (máximo 5 tópicos)
        3. Estrutura do documento (tipo de documento, seções principais)
        
        Texto para resumir:
        {text_content[:4000]}  # Limita para não exceder token limit
        """
        
        response = model.generate_content(
            prompt,
            safety_settings={
                HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
            }
        )
        
        return {
            "resumo": response.text.strip(),
            "palavras_chave": extract_keywords_from_text(text_content)
        }
        
    except Exception as e:
        return {"erro": f"Erro ao gerar resumo: {str(e)}"}

def extract_keywords_from_text(text):
    """
    Extrai palavras-chave importantes do texto
    """
    words = text.lower().split()
    
    stop_words = {
        'o', 'a', 'os', 'as', 'um', 'uma', 'uns', 'umas', 'de', 'do', 'da', 'dos', 'das',
        'em', 'no', 'na', 'nos', 'nas', 'por', 'para', 'com', 'sem', 'sob', 'sobre',
        'e', 'ou', 'mas', 'se', 'que', 'quando', 'como', 'onde', 'porque', 'então',
        'ele', 'ela', 'eles', 'elas', 'eu', 'tu', 'você', 'nós', 'vós', 'vocês',
        'este', 'esta', 'estes', 'estas', 'esse', 'essa', 'esses', 'essas',
        'aquele', 'aquela', 'aqueles', 'aquelas', 'isto', 'isso', 'aquilo'
    }
    
    keywords = []
    word_count = {}
    
    for word in words:
        clean_word = ''.join(c for c in word if c.isalnum())
        if len(clean_word) > 3 and clean_word not in stop_words:
            word_count[clean_word] = word_count.get(clean_word, 0) + 1
    
    sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)
    keywords = [word for word, count in sorted_words[:10]]
    
    return keywords

def search_text_in_document(text_content, search_term):
    """
    Busca palavras-chave no documento e retorna contexto
    """
    lines = text_content.split('\n')
    results = []
    
    for i, line in enumerate(lines):
        if search_term.lower() in line.lower():
            context_start = max(0, i - 1)
            context_end = min(len(lines), i + 2)
            context = '\n'.join(lines[context_start:context_end])
            
            results.append({
                "linha": i + 1,
                "texto": line.strip(),
                "contexto": context,
                "posicao": line.lower().find(search_term.lower())
            })
    
    return {
        "termo_busca": search_term,
        "total_ocorrencias": len(results),
        "resultados": results[:10]
    }


def extract_images_from_pdf(pdf_doc):
    """
    Extrai todas as imagens de um PDF
    Retorna lista com PIL Images e metadados
    """
    images = []
    
    for page_num in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_num)
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Filtrar imagens muito pequenas (ícones/decorações)
                if len(image_bytes) > 5000:  # > 5KB
                    pil_image = Image.open(io.BytesIO(image_bytes))
                    
                    # Filtrar por dimensões também
                    if pil_image.width > 100 and pil_image.height > 100:
                        images.append({
                            'page': page_num + 1,
                            'image': pil_image,
                            'size': len(image_bytes),
                            'dimensions': f"{pil_image.width}x{pil_image.height}"
                        })
            except Exception as e:
                print(f"Erro ao extrair imagem da página {page_num + 1}: {e}")
                continue
    
    return images


def extract_images_from_docx(file_content):
    """
    Extrai todas as imagens de um DOCX
    Retorna lista com PIL Images e metadados
    """
    doc = Document(io.BytesIO(file_content))
    images = []
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_bytes = rel.target_part.blob
                
                if len(image_bytes) > 5000:  # > 5KB
                    pil_image = Image.open(io.BytesIO(image_bytes))
                    
                    if pil_image.width > 100 and pil_image.height > 100:
                        images.append({
                            'image': pil_image,
                            'size': len(image_bytes),
                            'dimensions': f"{pil_image.width}x{pil_image.height}"
                        })
            except Exception as e:
                print(f"Erro ao extrair imagem do DOCX: {e}")
                continue
    
    return images


def analyze_image_with_gemini(pil_image, page_num=None):
    """
    Extrai texto da imagem com Tesseract OCR e analisa o conteúdo
    """
    try:
        # 1. Extrair texto com OCR
        texto_ocr = ""
        if TESSERACT_AVAILABLE:
            try:
                # Configurar para português
                texto_ocr = pytesseract.image_to_string(
                    pil_image,
                    lang='por',
                    config='--oem 3 --psm 6'
                ).strip()
            except Exception as e:
                print(f"Erro no OCR: {e}")
                texto_ocr = ""
        
        # 2. Se não tem texto, retornar indicação
        if not texto_ocr or len(texto_ocr) < 10:
            return "[Imagem sem texto legível detectado]", texto_ocr
        
        # 3. Usar Gemini para resumir/analisar o texto extraído
        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = f"""Você está analisando texto extraído de uma imagem dentro de um documento.

Texto extraído (OCR):
{texto_ocr}

Descreva em 2-3 linhas:
1. Que tipo de elemento visual é (banner, gráfico, tabela, etc)
def process_document_images(file_content, is_pdf, pdf_doc=None):
    """
    Processa todas as imagens do documento
    Extrai texto com OCR e analisa com Gemini
    """
    try:
        if not TESSERACT_AVAILABLE:
            print("⚠️ Tesseract não disponível, pulando análise de imagens")
            return []
        
        print("📸 Extraindo imagens do documento...")
        
        if is_pdf and pdf_doc:
            images = extract_images_from_pdf(pdf_doc)
        else:
            images = extract_images_from_docx(file_content)
        
        if not images:
            print("Nenhuma imagem encontrada no documento")
            return []
        
        print(f"🔍 {len(images)} imagens encontradas, processando com OCR...")
        
        image_descriptions = []
        for idx, img_info in enumerate(images, 1):
            try:
                print(f"  Processando imagem {idx}/{len(images)}...")
                
                # OCR + Análise com Gemini
                description, texto_ocr = analyze_image_with_gemini(
                    img_info['image'],
                    img_info.get('page')
                )
                
                if texto_ocr and len(texto_ocr) > 10:  # Só adicionar se tiver texto relevante
                    image_descriptions.append({
                        'numero': idx,
                        'pagina': img_info.get('page', 'N/A'),
                        'texto_ocr': texto_ocr,
                        'descricao': description,
                        'tamanho': img_info['size'],
                        'dimensoes': img_info['dimensions']
                    })
                
            except Exception as e:
                print(f"⚠️ Erro ao processar imagem {idx}: {e}")
                continue
        
        print(f"✅ {len(image_descriptions)} imagens com texto analisadas")
        return image_descriptions
        
    except Exception as e:
        print(f"❌ Erro ao processar imagens: {e}")
        return []   img_info['image'],
                    img_info.get('page')
                )
                
                image_descriptions.append({
                    'numero': idx,
                    'pagina': img_info.get('page', 'N/A'),
                    'descricao': description,
                    'tamanho': img_info['size'],
                    'dimensoes': img_info['dimensions']
                })
                
            except Exception as e:
                print(f"⚠️ Erro ao processar imagem {idx}: {e}")
                continue
        
        print(f"✅ {len(image_descriptions)} imagens analisadas com sucesso")
        return image_descriptions
        
    except Exception as e:
        print(f"❌ Erro ao processar imagens: {e}")
        return []

def process_document(file_content, file_name, file_type, gerar_resumo=True, analisar_imagens=True):
    """
    Processa documentos PDF e DOCX
    - PDFs: extração com PyMuPDF preservando estrutura
    - DOCX: extração com python-docx incluindo tabelas
    - Imagens: análise com Gemini Vision (opcional)
    - Resumo: geração automática com Gemini AI (opcional)
    """
    try:
        is_pdf = 'pdf' in file_type.lower() or file_name.lower().endswith('.pdf')
        is_docx = ('word' in file_type.lower() or 
                   'document' in file_type.lower() or 
                   file_name.lower().endswith(('.docx', '.doc')))
        
        if not (is_pdf or is_docx):
            return {
                "erro": "Tipo de arquivo não suportado. Use PDF ou DOCX.",
                "tipos_aceitos": ["PDF", "DOCX"]
            }
        
        print(f"Processando {'PDF' if is_pdf else 'DOCX'}: {file_name}")
        
        # Manter referência ao PDF doc para extrair imagens
        pdf_doc = None
        
        if is_pdf:
            pdf_doc = fitz.open(stream=file_content, filetype="pdf")
            # Extrair texto (reusar código existente inline)
            extracted_data = {
                "text_content": "",
                "structure": {
                    "pages": [],
                    "headings": [],
                    "paragraphs": [],
                    "tables": []
                },
                "metadata": {
                    "total_pages": len(pdf_doc),
                    "title": pdf_doc.metadata.get("title", ""),
                    "author": pdf_doc.metadata.get("author", "")
                }
            }
            
            full_text = ""
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                
                blocks = page.get_text("dict")
                page_text = ""
                page_structure = {
                    "page_number": page_num + 1,
                    "headings": [],
                    "paragraphs": [],
                    "images": []
                }
                
                for block in blocks["blocks"]:
                    if "lines" in block:
                        for line in block["lines"]:
                            line_text = ""
                            font_size = 0
                            
                            for span in line["spans"]:
                                line_text += span["text"]
                                font_size = max(font_size, span["size"])
                            
                            if line_text.strip():
                                if font_size > 14:
                                    page_structure["headings"].append({
                                        "text": line_text.strip(),
                                        "level": 1 if font_size > 18 else 2,
                                        "page": page_num + 1
                                    })
                                    extracted_data["structure"]["headings"].append({
                                        "text": line_text.strip(),
                                        "level": 1 if font_size > 18 else 2,
                                        "page": page_num + 1
                                    })
                                else:
                                    page_structure["paragraphs"].append(line_text.strip())
                                    extracted_data["structure"]["paragraphs"].append({
                                        "text": line_text.strip(),
                                        "page": page_num + 1
                                    })
                                
                                page_text += line_text + "\n"
                
                extracted_data["structure"]["pages"].append(page_structure)
                full_text += f"\n--- Página {page_num + 1} ---\n" + page_text
            
            extracted_data["text_content"] = full_text
            resultado = extracted_data
        else:
            resultado = extract_text_from_docx(file_content)
        
        if not resultado or 'erro' in resultado:
            if pdf_doc:
                pdf_doc.close()
            if image_descriptions:
                # Adicionar descrições das imagens ao texto
                resultado['text_content'] += "\n\n" + "="*60
                resultado['text_content'] += "\n📸 TEXTO EXTRAÍDO DE IMAGENS (OCR)\n"
                resultado['text_content'] += "="*60 + "\n"
                
                for desc in image_descriptions:
                    resultado['text_content'] += f"\n[Imagem {desc['numero']}"
                    if desc['pagina'] != 'N/A':
                        resultado['text_content'] += f" - Página {desc['pagina']}"
                    resultado['text_content'] += f" | {desc['dimensoes']}]\n"
                    resultado['text_content'] += f"💬 Texto: {desc['texto_ocr'][:300]}\n"
                    resultado['text_content'] += f"📝 Análise: {desc['descricao']}\n"
                
                resultado['imagens_analisadas'] = len(image_descriptions)
                    resultado['text_content'] += f"\n[Imagem {desc['numero']}"
                    if desc['pagina'] != 'N/A':
                        resultado['text_content'] += f" - Página {desc['pagina']}"
                    resultado['text_content'] += f" | {desc['dimensoes']}]\n"
                    resultado['text_content'] += f"📝 {desc['descricao']}\n"
                
                resultado['imagens_analisadas'] = len(image_descriptions)
        
        # Fechar PDF se foi aberto
        if pdf_doc:
            pdf_doc.close()
        
        resultado['arquivo_info'] = {
            'nome': file_name,
            'tipo': file_type,
            'tamanho_bytes': len(file_content),
            'formato': 'PDF' if is_pdf else 'DOCX',
            'total_imagens': len(image_descriptions)
        }
        
        if gerar_resumo and resultado.get('text_content'):
            print("Gerando resumo com Gemini...")
            resumo_data = generate_document_summary(resultado['text_content'])
            
            if 'erro' not in resumo_data:
                resultado['resumo'] = resumo_data.get('resumo')
                resultado['palavras_chave'] = resumo_data.get('palavras_chave')
        
        print("Documento processado com sucesso!")
        return resultado
        
    except Exception as e:
        print(f"Erro: {str(e)}")
        if pdf_doc:
            pdf_doc.close()
        return {"erro": f"Erro ao processar documento: {str(e)}"}
